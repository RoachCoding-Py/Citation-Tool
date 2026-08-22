"""
================================================================================
Stellenbosch Law Citation Pinpointer & Source Audit Tool (FREE GEMINI EDITION)
Developer: Aidan Roach
Degree/Institution: LLB Student, Stellenbosch University
================================================================================
"""

import os
import re
import io
from typing import List
import pandas as pd
import fitz  # PyMuPDF
import pdfplumber
import docx
from docx.shared import Pt, Inches, RGBColor
import streamlit as st
from pydantic import BaseModel, Field
from streamlit_local_storage import LocalStorage

# Safely load python-dotenv if installed
try:
    from dotenv import load_dotenv

    load_dotenv()
except ImportError:
    pass

# Import Google GenAI SDK
try:
    from google import genai
    from google.genai import types
except ImportError:
    st.error("Missing required library! Please run: `pip install google-genai` in your PyCharm terminal.")


# ==========================================
# 1. HELPER FUNCTIONS & CLEANING
# ==========================================

def clean_legal_text(text: str) -> str:
    """Removes common header/footer artifacts from JutaIQ, LexisNexis SA, and SAFLII."""
    text = re.sub(r"Downloaded from JutaIQ on \d{2}/\d{2}/\d{4}.*", "", text)
    text = re.sub(r"LexisNexis South Africa \([0-9-]+\).*", "", text)
    text = re.sub(r"SAFLII Note:.*", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def normalize_stellenbosch_citations(text: str) -> str:
    """
    Automated Rule-Based NLP Normalizer.
    Converts non-standard citations to Stellenbosch Faculty format:
    - 'Section 12' / 'Sec 12' -> 's 12'
    - 'Paragraph 10' / 'Par 10' -> 'para 10'
    """
    text = re.sub(r"\b(?:Section|Sec\.)\s+(\d+)", r"s \1", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(?:Sections|Secs\.)\s+(\d+)", r"ss \1", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(?:Paragraph|Par\.|para\.)\s+(\d+)", r"para \1", text, flags=re.IGNORECASE)
    text = re.sub(r"\b(?:Paragraphs|Pars\.|paras\.)\s+(\d+)", r"paras \1", text, flags=re.IGNORECASE)
    return text


def parse_docx_file(file_bytes: bytes) -> str:
    """Extracts text from Word documents."""
    doc = docx.Document(io.BytesIO(file_bytes))
    return "\n\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])


def parse_pdf_draft(file_bytes: bytes) -> str:
    """Extracts plain text from draft PDFs."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text = [page.get_text("text").strip() for page in doc if page.get_text("text").strip()]
    doc.close()
    return clean_legal_text("\n\n".join(pages_text))


def parse_and_anchor_pdf(file_bytes: bytes) -> str:
    """Extracts PDF text and injects structural [[PAGE X]] markers."""
    anchored_text = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num in range(len(doc)):
            page_text = clean_legal_text(doc[page_num].get_text("text"))
            anchored_text.append(f"[[PAGE {page_num + 1}]]\n{page_text}")
        doc.close()
    except Exception:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = clean_legal_text(page.extract_text() or "")
                anchored_text.append(f"[[PAGE {page_num + 1}]]\n{page_text}")

    full_text = "\n\n".join(anchored_text)
    return re.sub(r"\n\[(\d+)\]\s*", r"\n[[PARA \1]] ", full_text)


def generate_docx_download(annotated_text: str, citation_style: str) -> bytes:
    """Generates an academic-formatted .docx file."""
    doc = docx.Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    title = doc.add_heading(level=0)
    run = title.add_run("Annotated Legal Draft & Source Audit")
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 33, 71)

    subtitle = doc.add_paragraph()
    sub_run = subtitle.add_run(f"Faculty Style: {citation_style} | Format: SU Footnote System")
    sub_run.font.name = "Times New Roman"
    sub_run.font.italic = True
    sub_run.font.size = Pt(10)

    for p_text in annotated_text.split("\n\n"):
        if p_text.strip():
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(p_text.strip())
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# ==========================================
# 2. SCHEMA & GEMINI INFERENCE ENGINE
# ==========================================

STELLENBOSCH_CITATION_SYSTEM_PROMPT = """You are a senior South African Legal Tech Specialist at Stellenbosch University Faculty of Law.
Your role is to scan a student's legal draft against provided source material, identify claims wrapped in quotation marks ("..."), match them to exact source locations, and generate formatted pinpoint citations.

FORMATTING RULES:
1. Cases: *S v Makwanyane* 1995 (3) SA 391 (CC) at para 100
2. Statutes: s 12(1)(a) of the Constitution of the Republic of South Africa, 1996
3. Insert numeric markers like [^1] immediately after quoted claims, and append a Footnotes section at the very end of the document.
4. If a claim cannot be matched with high confidence (>85%), mark pinpoint as '[UNVERIFIED - PINPOINT NOT FOUND]' and confidence_status as 'UNVERIFIED'.
"""


class AuditItem(BaseModel):
    draft_claim: str = Field(description="The exact claim or quoted text from the student's draft.")
    matched_source_title: str = Field(description="Title/label of the matching source document.")
    pinpoint_citation: str = Field(description="Formatted pinpoint citation, or '[UNVERIFIED - PINPOINT NOT FOUND]'.")
    verbatim_source_excerpt: str = Field(description="Verbatim excerpt from source containing the anchor.")
    confidence_status: str = Field(description="'VERIFIED' if confidence > 85%, otherwise 'UNVERIFIED'.")


class AuditReport(BaseModel):
    annotated_draft: str = Field(
        description="The draft with footnote markers [^1] after quotes and a Footnotes list appended at the end.")
    audit_items: List[AuditItem] = Field(description="Structured verification items.")


def run_citation_audit_gemini(
        draft_text: str,
        sources: List[dict],
        citation_style: str,
        gemini_key: str
) -> AuditReport:
    client = genai.Client(api_key=gemini_key)

    formatted_sources = "\n\n".join([f"=== SOURCE: {s['title']} ===\n{s['content']}" for s in sources])
    user_prompt = f"SELECTED STYLE: {citation_style}\n\nDRAFT:\n{draft_text}\n\nSOURCES:\n{formatted_sources}"

    response = client.models.generate_content(
        model='gemini-3.6-flash',
        contents=user_prompt,
        config=types.GenerateContentConfig(
            system_instruction=STELLENBOSCH_CITATION_SYSTEM_PROMPT,
            temperature=0.0,
            response_mime_type="application/json",
            response_schema=AuditReport,
        ),
    )

    return AuditReport.model_validate_json(response.text)

# ==========================================
# 3. STREAMLIT INTERFACE
# ==========================================

def main():
    st.set_page_config(page_title="SU Law Citation Pinpointer", page_icon="⚖️", layout="wide")

    # Initialize Local Storage
    local_storage = LocalStorage()

    st.title("⚖️ Stellenbosch Law Citation Pinpointer & Audit Tool")
    st.caption("Created by Aidan Roach | Faculty of Law, Stellenbosch University (Free Gemini Edition)")
    st.info("🔔 This tool runs on a free-tier API. If it stops responding, the daily free usage limit has likely been reached — please try again tomorrow.")
    # --- SIDEBAR ---
    st.sidebar.header("⚙️ Configuration")
    citation_style = st.sidebar.selectbox(
        "Faculty Citation Style",
        options=["Stellenbosch Law Faculty Guidelines (Footnotes)", "OSCOLA / SA Law Journal Style"]
    )
    auto_normalize = st.sidebar.checkbox("Auto-Normalize Citation Formatting", value=True)

    st.sidebar.markdown("---")
    st.sidebar.subheader("🔑 API Key")

    env_key = os.getenv("GEMINI_API_KEY", "")
    gemini_api_key = st.sidebar.text_input("Gemini API Key (Free)", value=env_key, type="password")

    with st.sidebar.expander("📖 SU Citation Cheat Sheet"):
        st.markdown("""
        **Cases:**  
        `*S v Makwanyane* 1995 (3) SA 391 (CC) at para 100`  

        **Statutes:**  
        `s 12(1)(a) of the Constitution...` *(lowercase 's')*  
        `ss 14-16 of the Companies Act 71 of 2008`
        """)

    # --- BANNERS ---
    st.warning(
        "📌 **Quotation Requirement:** Enclose every statement or claim you want pinpointed inside **quotation marks** (e.g., *\"the right to dignity is non-derogable\"*).")
    st.info(
        "ℹ️ **SU Faculty Notice:** Enforces the mandatory **footnote referencing system** (`[^1]`). In-text parenthetical citations are not permitted under SU guidelines.")

    # --- MAIN COLUMNS ---
    col1, col2 = st.columns([1, 1], gap="large")

    draft_input_text = ""
    with col1:
        st.subheader("1. Student Legal Draft")
        draft_mode = st.radio("Input Method:", ["Paste Text Directly", "Upload File (.docx / .pdf)"], horizontal=True)

        if draft_mode == "Paste Text Directly":
            # Retrieve saved draft from browser storage if present
            saved_draft = local_storage.getItem("user_draft_text") or ""

            draft_input_text = st.text_area(
                "Paste assignment here:",
                value=saved_draft,
                height=300,
                key="draft_text_input"
            )

            # Save to browser storage whenever text changes
            if draft_input_text:
                local_storage.setItem("user_draft_text", draft_input_text)

        else:
            uploaded_draft = st.file_uploader("Upload Word or PDF:", type=["docx", "pdf"])
            if uploaded_draft:
                file_bytes = uploaded_draft.read()
                draft_input_text = parse_docx_file(file_bytes) if uploaded_draft.name.endswith(
                    ".docx") else parse_pdf_draft(file_bytes)
                st.success(f"Loaded '{uploaded_draft.name}'")

    with col2:
        st.subheader("2. Reference Source Materials")
        num_sources = st.number_input("Number of Sources", min_value=1, max_value=5, value=2)
        sources_payload = []

        for i in range(int(num_sources)):
            st.markdown(f"**Source #{i + 1}**")
            src_title = st.text_input(f"Source Title #{i + 1}", value=f"Source {i + 1}", key=f"t_{i}")
            src_file = st.file_uploader(f"Upload Source PDF #{i + 1}", type=["pdf"], key=f"f_{i}")

            if src_file:
                parsed_text = parse_and_anchor_pdf(src_file.read())
                sources_payload.append({"title": src_title, "content": parsed_text})
                st.success(f"Parsed {src_file.name}")

    st.markdown("---")

    # --- ACTION BUTTON ---
    if st.button("🔍 Scan & Generate Pinpoints", type="primary", use_container_width=True):
        if not draft_input_text.strip():
            st.error("Please provide a draft.")
            return
        if not sources_payload:
            st.error("Please upload at least one source PDF.")
            return
        if not gemini_api_key:
            st.error("Missing Gemini API key! Paste it in the sidebar or save it in your .env file.")
            return

        if auto_normalize:
            draft_input_text = normalize_stellenbosch_citations(draft_input_text)

        with st.spinner("Scanning claims with Gemini 2.5 Flash..."):
            try:
                report = run_citation_audit_gemini(
                    draft_text=draft_input_text,
                    sources=sources_payload,
                    citation_style=citation_style,
                    gemini_key=gemini_api_key
                )
                st.session_state["audit_report"] = report
                st.success("Scan Complete!")
            except Exception as e:
                st.error(f"Execution Error: {str(e)}")

    # --- RESULTS ---
    if "audit_report" in st.session_state:
        report: AuditReport = st.session_state["audit_report"]

        st.markdown("## 📊 Audit Results")
        total_claims = len(report.audit_items)
        verified_claims = sum(1 for item in report.audit_items if item.confidence_status == "VERIFIED")
        unverified_claims = total_claims - verified_claims

        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Quoted Claims", total_claims)
        m2.metric("Verified Pinpoints", verified_claims)
        m3.metric("Unverified Flags", unverified_claims)
        m4.metric("Model", "Gemini 2.5 Flash (Free)")

        tab1, tab2 = st.tabs(["📝 Annotated Draft", "🔎 Audit Table"])

        with tab1:
            st.text_area("Annotated Document", value=report.annotated_draft, height=350)
            docx_bytes = generate_docx_download(report.annotated_draft, citation_style)
            st.download_button("📄 Download .docx", data=docx_bytes, file_name="annotated_draft.docx",
                               use_container_width=True)

        with tab2:
            table_data = [{
                "Status": "✅ VERIFIED" if item.confidence_status == "VERIFIED" else "⚠️ UNVERIFIED",
                "Draft Claim": item.draft_claim,
                "Source": item.matched_source_title,
                "Pinpoint": item.pinpoint_citation,
                "Verbatim Excerpt": item.verbatim_source_excerpt
            } for item in report.audit_items]
            st.dataframe(pd.DataFrame(table_data), use_container_width=True)


if __name__ == "__main__":
    main()